from __future__ import annotations

"""Build editable Word/Excel/PDF source ledgers for the competition demo."""

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "演示业务源文件"
ROSTER_FILE = BASE_DIR / "roster_students.json"
SOURCE_NOTE = "比赛演示模拟数据：用于功能验证、录屏和答辩展示，不对应真实个人。"


def roster_by_id() -> dict[str, dict[str, str]]:
    return {str(item["student_id"]): item for item in load_roster()}


def load_roster() -> list[dict[str, str]]:
    return json.loads(ROSTER_FILE.read_text(encoding="utf-8"))


def student_row(student_id: str, dorm: str) -> dict[str, str]:
    student = roster_by_id()[student_id]
    return {
        "学号": student_id,
        "姓名": student["name"],
        "性别": student["gender"],
        "脱敏手机号": student["phone_masked"],
        "学院": student["college"],
        "专业": student["major"],
        "年级": student["grade"],
        "班级": student["class_name"],
        "宿舍": dorm,
    }


def write_xlsx(path: Path, sheet_name: str, headers: list[str], rows: list[dict[str, str]]) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = sheet_name
    title = f"浙江经济职业技术学院 {sheet_name}（比赛演示）"
    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    cell = sheet.cell(1, 1, title)
    cell.font = Font(size=14, bold=True, color="FFFFFF")
    cell.fill = PatternFill("solid", fgColor="1F4E78")
    cell.alignment = Alignment(horizontal="center", vertical="center")
    sheet.row_dimensions[1].height = 28
    sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
    note = sheet.cell(2, 1, SOURCE_NOTE)
    note.font = Font(size=10, color="666666")
    note.alignment = Alignment(horizontal="left", vertical="center")
    sheet.row_dimensions[2].height = 22
    for index, header in enumerate(headers, start=1):
        cell = sheet.cell(3, index, header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2F75B5")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row_index, row in enumerate(rows, start=4):
        for col_index, header in enumerate(headers, start=1):
            cell = sheet.cell(row_index, col_index, row.get(header, ""))
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            if row_index % 2 == 0:
                cell.fill = PatternFill("solid", fgColor="EAF2F8")
    sheet.freeze_panes = "A4"
    sheet.auto_filter.ref = f"A3:{get_column_letter(len(headers))}{len(rows) + 3}"
    for index, header in enumerate(headers, start=1):
        width = max(len(header) * 2.2, 12)
        if header in ("备注", "建议说明"):
            width = 40
        if header in ("班级", "专业"):
            width = 18
        sheet.column_dimensions[get_column_letter(index)].width = width
    workbook.save(path)


def build_roster() -> list[dict[str, str]]:
    dorm_prefix = {"S604124移动": "8-", "P603124数媒": "5-", "P603223数媒": "5-", "S603323数媒": "6-", "W602325网络": "7-", "W602425网络": "7-"}
    rows = []
    for index, student in enumerate(load_roster(), start=1):
        dorm = f"{dorm_prefix[student['class_name']]}{300 + index % 18}"
        rows.append(student_row(str(student["student_id"]), dorm))
    return rows


def build_documents() -> list[Path]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    roster = roster_by_id()
    students = [
        "60412403", "60412408", "60412415", "60312401", "60322304", "60332305", "60232502", "60242504"
    ]
    roster_rows = build_roster()
    roster_file = OUTPUT_DIR / "01_学生基础名册_比赛演示.xlsx"
    write_xlsx(roster_file, "学生名册", list(roster_rows[0]), roster_rows)

    attendance_rows = [
        {"记录编号": "ATT-202608-001", "统计周期": "2026-07-16至2026-08-14", "学号": "60412403", "姓名": roster["60412403"]["name"], "班级": roster["60412403"]["class_name"], "缺勤次数": "5", "迟到次数": "2", "学业趋势": "明显下降", "不及格课程数": "2", "资助状态": "家庭经济困难认定在档", "备注": "连续一周出现3次旷课，建议班主任与辅导员共同核实。"},
        {"记录编号": "ATT-202608-002", "统计周期": "2026-07-16至2026-08-14", "学号": "60412408", "姓名": roster["60412408"]["name"], "班级": roster["60412408"]["class_name"], "缺勤次数": "3", "迟到次数": "4", "学业趋势": "下降", "不及格课程数": "1", "资助状态": "一般", "备注": "需安排学习计划沟通。"},
        {"记录编号": "ATT-202608-003", "统计周期": "2026-07-16至2026-08-14", "学号": "60412415", "姓名": roster["60412415"]["name"], "班级": roster["60412415"]["class_name"], "缺勤次数": "1", "迟到次数": "1", "学业趋势": "稳定", "不及格课程数": "0", "资助状态": "一般", "备注": "常规关注。"},
        {"记录编号": "ATT-202608-004", "统计周期": "2026-07-16至2026-08-14", "学号": "60312401", "姓名": roster["60312401"]["name"], "班级": roster["60312401"]["class_name"], "缺勤次数": "2", "迟到次数": "3", "学业趋势": "下降", "不及格课程数": "1", "资助状态": "一般", "备注": "建议联系任课教师了解课堂表现。"},
        {"记录编号": "ATT-202608-005", "统计周期": "2026-07-16至2026-08-14", "学号": "60322304", "姓名": roster["60322304"]["name"], "班级": roster["60322304"]["class_name"], "缺勤次数": "0", "迟到次数": "3", "学业趋势": "稳定", "不及格课程数": "0", "资助状态": "一般", "备注": "常规提醒。"},
        {"记录编号": "ATT-202608-006", "统计周期": "2026-07-16至2026-08-14", "学号": "60332305", "姓名": roster["60332305"]["name"], "班级": roster["60332305"]["class_name"], "缺勤次数": "4", "迟到次数": "1", "学业趋势": "下降", "不及格课程数": "2", "资助状态": "待认定", "备注": "建议同步核实资助和学业支持需求。"},
        {"记录编号": "ATT-202608-007", "统计周期": "2026-07-16至2026-08-14", "学号": "60232502", "姓名": roster["60232502"]["name"], "班级": roster["60232502"]["class_name"], "缺勤次数": "0", "迟到次数": "1", "学业趋势": "稳定", "不及格课程数": "0", "资助状态": "一般", "备注": "常规关注。"},
        {"记录编号": "ATT-202608-008", "统计周期": "2026-07-16至2026-08-14", "学号": "60242504", "姓名": roster["60242504"]["name"], "班级": roster["60242504"]["class_name"], "缺勤次数": "2", "迟到次数": "2", "学业趋势": "稳定", "不及格课程数": "0", "资助状态": "一般", "备注": "建议关注实训参与情况。"},
    ]
    attendance_file = OUTPUT_DIR / "02_考勤与学业汇总_2026年8月.xlsx"
    write_xlsx(attendance_file, "考勤与学业", list(attendance_rows[0]), attendance_rows)

    training_rows = [
        {"记录编号": "TRN-202608-001", "记录日期": "2026-08-14", "学号": "60412403", "姓名": roster["60412403"]["name"], "班级": roster["60412403"]["class_name"], "实训项目": "移动应用综合实训", "缺交日志数": "3", "操作异常数": "1", "实训指导教师": "周老师", "记录说明": "日志缺交，接口联调环节未按规范提交。"},
        {"记录编号": "TRN-202608-002", "记录日期": "2026-08-14", "学号": "60412408", "姓名": roster["60412408"]["name"], "班级": roster["60412408"]["class_name"], "实训项目": "移动应用综合实训", "缺交日志数": "1", "操作异常数": "2", "实训指导教师": "周老师", "记录说明": "两次未完成代码提交规范检查。"},
        {"记录编号": "TRN-202608-003", "记录日期": "2026-08-13", "学号": "60312401", "姓名": roster["60312401"]["name"], "班级": roster["60312401"]["class_name"], "实训项目": "短视频创作实训", "缺交日志数": "2", "操作异常数": "1", "实训指导教师": "李老师", "记录说明": "缺少阶段性复盘日志。"},
        {"记录编号": "TRN-202608-004", "记录日期": "2026-08-13", "学号": "60332305", "姓名": roster["60332305"]["name"], "班级": roster["60332305"]["class_name"], "实训项目": "数字内容制作实训", "缺交日志数": "2", "操作异常数": "2", "实训指导教师": "胡老师", "记录说明": "操作规范提醒两次，需安排补训。"},
        {"记录编号": "TRN-202608-005", "记录日期": "2026-08-12", "学号": "60242504", "姓名": roster["60242504"]["name"], "班级": roster["60242504"]["class_name"], "实训项目": "网络综合布线实训", "缺交日志数": "1", "操作异常数": "0", "实训指导教师": "陈老师", "记录说明": "日志待补交。"},
    ]
    training_file = OUTPUT_DIR / "03_实训日志与操作记录_2026年8月.xlsx"
    write_xlsx(training_file, "实训记录", list(training_rows[0]), training_rows)

    screening_rows = [
        {"记录编号": "PSY-202608-001", "筛查日期": "2026-08-10", "学号": "60412403", "姓名": roster["60412403"]["name"], "班级": roster["60412403"]["class_name"], "筛查建议": "建议关怀跟进", "建议说明": "问卷显示近期压力感较高，建议由辅导员了解学习适应与生活支持需求；非诊断结论。"},
        {"记录编号": "PSY-202608-002", "筛查日期": "2026-08-10", "学号": "60332305", "姓名": roster["60332305"]["name"], "班级": roster["60332305"]["class_name"], "筛查建议": "建议关怀跟进", "建议说明": "建议安排一次关怀沟通，了解作息和课程负担；非诊断结论。"},
        {"记录编号": "PSY-202608-003", "筛查日期": "2026-08-10", "学号": "60412408", "姓名": roster["60412408"]["name"], "班级": roster["60412408"]["class_name"], "筛查建议": "常规关注", "建议说明": "未见需专项跟进信号。"},
    ]
    screening_file = OUTPUT_DIR / "04_心理筛查结果_2026年秋季.xlsx"
    write_xlsx(screening_file, "心理筛查", list(screening_rows[0]), screening_rows)

    talks = [
        {"记录编号": "FUP-20260518-001", "谈话日期": "2026-05-18", "学号": "60412403", "姓名": roster["60412403"]["name"], "班级": roster["60412403"]["class_name"], "谈话人": "林老师（辅导员）", "谈话主题": "学习适应与压力关怀", "谈话摘要": "学生反馈课程任务集中、时间安排欠合理，愿意在任课教师指导下补齐学习任务。", "下次行动": "2026-08-18 前复访，核实补考复习、实训日志补交与生活适应情况。", "下次跟进日期": "2026-08-18"},
        {"记录编号": "FUP-20260812-002", "谈话日期": "2026-08-12", "学号": "60332305", "姓名": roster["60332305"]["name"], "班级": roster["60332305"]["class_name"], "谈话人": "林老师（辅导员）", "谈话主题": "学业与资助需求核实", "谈话摘要": "学生表示近期存在课程理解困难，已说明可申请学业辅导和资助咨询。", "下次行动": "2026-09-02 前确认补训安排和资助认定材料。", "下次跟进日期": "2026-09-02"},
        {"记录编号": "FUP-20260813-003", "谈话日期": "2026-08-13", "学号": "60412408", "姓名": roster["60412408"]["name"], "班级": roster["60412408"]["class_name"], "谈话人": "赵老师（班主任）", "谈话主题": "实训纪律与学习计划", "谈话摘要": "已沟通实训提交规范，学生承诺在本周内补齐日志。", "下次行动": "2026-08-20 前由班主任核对日志补交情况。", "下次跟进日期": "2026-08-20"},
    ]
    talk_file = OUTPUT_DIR / "05_辅导员谈心谈话记录_2026年8月.docx"
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("浙江经济职业技术学院\n学生谈心谈话记录（比赛演示）")
    run.bold = True
    run.font.size = Pt(16)
    note = document.add_paragraph(SOURCE_NOTE)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].font.size = Pt(9)
    headers = list(talks[0])
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for talk in talks:
        cells = table.add_row().cells
        for index, header in enumerate(headers):
            cells[index].text = talk[header]
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)
    document.add_paragraph("说明：筛查和谈话信息仅用于比赛演示中的关怀、提醒和任务闭环，不作心理诊断或处分依据。")
    document.save(talk_file)

    task_rows = [
        ["SUP-202608-001", "60412403", roster["60412403"]["name"], roster["60412403"]["class_name"], "林老师（辅导员）", "2026-08-18", "跟进中", "紧急", "完成学习状态与实训日志核实", "复访并记录补考计划、日志补交情况"],
        ["SUP-202608-002", "60412408", roster["60412408"]["name"], roster["60412408"]["class_name"], "赵老师（班主任）", "2026-08-20", "待处理", "重点", "完成实训纪律与日志补交提醒", "联系实训指导教师核验补交结果"],
        ["SUP-202608-003", "60332305", roster["60332305"]["name"], roster["60332305"]["class_name"], "林老师（辅导员）", "2026-09-02", "跟进中", "重点", "落实学业辅导和资助材料核实", "确认补训安排与资助认定进度"],
    ]
    task_file = OUTPUT_DIR / "06_学生帮扶任务清单_2026年8月.pdf"
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    style = ParagraphStyle("cn", parent=styles["Normal"], fontName="STSong-Light", fontSize=8, leading=12)
    heading = ParagraphStyle("heading", parent=style, fontSize=16, leading=22, alignment=1)
    doc = SimpleDocTemplate(str(task_file), pagesize=A4, rightMargin=1.2 * cm, leftMargin=1.2 * cm, topMargin=1.4 * cm, bottomMargin=1.4 * cm)
    flow = [Paragraph("浙江经济职业技术学院 学生帮扶任务清单（比赛演示）", heading), Spacer(1, 0.25 * cm), Paragraph(SOURCE_NOTE, style), Spacer(1, 0.25 * cm)]
    headers = ["任务编号", "学生", "班级", "负责人", "期限", "状态", "优先级", "任务目标"]
    display_rows = [headers]
    for row in task_rows:
        display_rows.append([row[0], f"{row[2]}\n{row[1]}", row[3], row[4], row[5], row[6], row[7], row[8]])
    table = Table(display_rows, colWidths=[2.1*cm, 2.2*cm, 2.3*cm, 2.4*cm, 1.6*cm, 1.3*cm, 1.2*cm, 5.0*cm])
    table.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), "STSong-Light"), ("FONTSIZE", (0,0), (-1,-1), 7.3),
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#2F75B5")), ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#9EADBB")), ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("BACKGROUND", (0,1), (-1,-1), colors.HexColor("#F6FAFE")), ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4),
        ("TOPPADDING", (0,0), (-1,-1), 6), ("BOTTOMPADDING", (0,0), (-1,-1), 6),
    ]))
    flow += [table, Spacer(1, 0.35 * cm), Paragraph("系统导入校验记录（供程序读取，不对外展示）", style)]
    for row in task_rows:
        flow.append(Paragraph("TASK|" + "|".join(row), style))
    flow.append(Spacer(1, 0.2 * cm))
    flow.append(Paragraph("说明：帮扶任务应由授权人员结合正式业务系统复核，任何预警不构成对学生的标签或最终认定。", style))
    doc.build(flow)
    return [roster_file, attendance_file, training_file, screening_file, talk_file, task_file]


if __name__ == "__main__":
    for item in build_documents():
        print(item)
