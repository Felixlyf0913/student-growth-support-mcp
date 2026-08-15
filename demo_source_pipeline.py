from __future__ import annotations

"""Import competition-demo operational documents into the persistent record store.

The documents are deliberately ordinary office files: staff can maintain the
roster, attendance, training, screening, talk and task ledgers in their daily
workflow, while this module turns their structured fields into searchable
student profiles and reminders.
"""

import hashlib
import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from persistent_store import PersistentStore


SOURCE_VERSION = "2026-08-demo-v2"
SOURCE_LABEL = "比赛演示模拟业务台账"


def _text(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _number(value: Any) -> int:
    try:
        return int(float(value or 0))
    except (TypeError, ValueError):
        return 0


def _rows_from_workbook(
    path: Path,
    sheet_name: str,
    header_key: str = "学号",
) -> list[dict[str, str]]:
    workbook = load_workbook(path, data_only=True)
    if sheet_name not in workbook.sheetnames:
        raise ValueError(f"{path.name} 中未找到工作表：{sheet_name}")
    sheet = workbook[sheet_name]
    header_row = 0
    headers: list[str] = []
    # Daily ledgers often contain a title and a data-usage note before the
    # table. Locate the actual header instead of imposing a machine-only file.
    for row_index in range(1, min(sheet.max_row, 12) + 1):
        candidate = [_text(cell.value) for cell in sheet[row_index]]
        if header_key in candidate:
            header_row = row_index
            headers = candidate
            break
    if not header_row:
        raise ValueError(
            f"{path.name} 的工作表“{sheet_name}”未找到包含“{header_key}”的表头行"
        )
    rows: list[dict[str, str]] = []
    for values in sheet.iter_rows(min_row=header_row + 1, values_only=True):
        row = {headers[index]: _text(value) for index, value in enumerate(values) if headers[index]}
        if any(row.values()):
            rows.append(row)
    return rows


def _rows_from_docx(path: Path) -> list[dict[str, str]]:
    document = Document(path)
    if not document.tables:
        raise ValueError(f"{path.name} 中未找到谈话记录表格")
    table = document.tables[0]
    headers = [_text(cell.text) for cell in table.rows[0].cells]
    rows: list[dict[str, str]] = []
    for row_cells in table.rows[1:]:
        row = {
            headers[index]: _text(cell.text)
            for index, cell in enumerate(row_cells.cells)
            if index < len(headers) and headers[index]
        }
        if any(row.values()):
            rows.append(row)
    return rows


def _rows_from_pdf(path: Path) -> list[dict[str, str]]:
    text = "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
    rows: list[dict[str, str]] = []
    for line in text.splitlines():
        if not line.startswith("TASK|"):
            continue
        values = [value.strip() for value in line.split("|")]
        if len(values) != 11:
            raise ValueError(f"{path.name} 中的任务导入行字段数不正确：{line}")
        rows.append(
            {
                "task_id": values[1],
                "student_id": values[2],
                "student_name": values[3],
                "class_name": values[4],
                "owner": values[5],
                "due_date": values[6],
                "status": values[7],
                "priority": values[8],
                "objective": values[9],
                "next_action": values[10],
            }
        )
    if not rows:
        raise ValueError(f"{path.name} 中未找到可导入的 TASK 任务行")
    return rows


def _document_manifest(path: Path, record_count: int, document_type: str) -> dict[str, Any]:
    digest = hashlib.sha256(path.read_bytes()).hexdigest()[:16]
    return {
        "document_id": f"SRC-{digest}",
        "file_name": path.name,
        "file_type": path.suffix.lstrip(".").upper(),
        "document_type": document_type,
        "record_count": record_count,
        "source_label": SOURCE_LABEL,
        "checksum": digest,
        "imported_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "note": "比赛演示模拟数据，供功能验证、录屏和答辩展示使用。",
    }


def _base_profile(student: dict[str, str]) -> dict[str, Any]:
    return {
        "student_id": student["学号"],
        "name": student["姓名"],
        "gender": student["性别"],
        "phone_masked": student["脱敏手机号"],
        "college": student["学院"],
        "major": student["专业"],
        "grade": student["年级"],
        "class_name": student["班级"],
        "dorm": student.get("宿舍", "未登记"),
        "attendance_absences_30d": 0,
        "late_count_30d": 0,
        "gpa_trend": "稳定",
        "failed_courses": 0,
        "training_log_missing": 0,
        "training_issue_count": 0,
        "dorm_feedback": "正常",
        "financial_status": "一般",
        "talk_records": 0,
        "last_followup": "",
        "care_level": "低关注",
        "notes": "暂无明显异常，保持常规关注。",
        "data_label": SOURCE_LABEL,
        "data_source": "项目文件夹中的演示业务台账已导入 PostgreSQL",
    }


def _build_profiles(
    roster: list[dict[str, str]],
    attendance: list[dict[str, str]],
    training: list[dict[str, str]],
    screening: list[dict[str, str]],
    talks: list[dict[str, str]],
) -> list[dict[str, Any]]:
    profiles = {_text(row["学号"]): _base_profile(row) for row in roster}
    for row in attendance:
        profile = profiles.get(row["学号"])
        if not profile:
            continue
        profile["attendance_absences_30d"] = _number(row.get("缺勤次数"))
        profile["late_count_30d"] = _number(row.get("迟到次数"))
        profile["gpa_trend"] = row.get("学业趋势") or "稳定"
        profile["failed_courses"] = _number(row.get("不及格课程数"))
        profile["financial_status"] = row.get("资助状态") or "一般"
    for row in training:
        profile = profiles.get(row["学号"])
        if not profile:
            continue
        profile["training_log_missing"] = _number(row.get("缺交日志数"))
        profile["training_issue_count"] = _number(row.get("操作异常数"))
    for row in screening:
        profile = profiles.get(row["学号"])
        if not profile:
            continue
        result = row.get("筛查建议", "常规关注")
        if result != "常规关注":
            profile["dorm_feedback"] = f"筛查建议：{result}（非诊断）"
            profile["notes"] = row.get("建议说明") or "建议由辅导员开展关怀沟通并持续观察。"
    for row in talks:
        profile = profiles.get(row["学号"])
        if not profile:
            continue
        profile["talk_records"] += 1
        date = row.get("谈话日期", "")
        if date and date >= profile["last_followup"]:
            profile["last_followup"] = date
        if row.get("谈话摘要"):
            profile["notes"] = row["谈话摘要"]
    return list(profiles.values())


def _normalize_roster(rows: list[dict[str, str]]) -> list[dict[str, Any]]:
    return [
        {
            "student_id": row["学号"],
            "name": row["姓名"],
            "gender": row["性别"],
            "phone_masked": row["脱敏手机号"],
            "college": row["学院"],
            "major": row["专业"],
            "grade": row["年级"],
            "class_name": row["班级"],
            "dorm": row.get("宿舍", "未登记"),
            "source": "01_学生基础名册_比赛演示.xlsx",
            "data_label": SOURCE_LABEL,
        }
        for row in rows
    ]


def _normalize_followups(rows: list[dict[str, str]]) -> list[dict[str, Any]]:
    return [
        {
            "record_id": row["记录编号"],
            "student_id": row["学号"],
            "student_name": row["姓名"],
            "class_name": row["班级"],
            "created_at": f"{row['谈话日期']} 14:00",
            "owner": row["谈话人"],
            "summary": row["谈话摘要"],
            "next_action": row["下次行动"],
            "next_followup_date": row["下次跟进日期"],
            "status": "跟进中",
            "source": "05_辅导员谈心谈话记录_2026年8月.docx",
            "data_label": SOURCE_LABEL,
        }
        for row in rows
    ]


def _normalize_tasks(rows: list[dict[str, str]], profiles: list[dict[str, Any]]) -> list[dict[str, Any]]:
    profile_by_id = {row["student_id"]: row for row in profiles}
    normalized: list[dict[str, Any]] = []
    for row in rows:
        profile = profile_by_id.get(row["student_id"], {})
        normalized.append(
            {
                **row,
                "attention_score": 0,
                "attention_level": "待核实",
                "risk_reasons": [],
                "measures": [row["next_action"]],
                "created_by": "比赛演示数据导入",
                "created_at": "2026-08-15 09:00",
                "updated_at": "2026-08-15 09:00",
                "completed_at": "",
                "progress_history": [],
                "notification": {"requested": False, "sent": False, "channel": ""},
                "data_label": SOURCE_LABEL,
                "student_major": profile.get("major", ""),
            }
        )
    return normalized


def _normalize_training_rooms(rows: list[dict[str, str]]) -> list[dict[str, Any]]:
    return [
        {
            "room_id": row["实训室编号"],
            "name": row["实训室名称"],
            "building": row["所在楼宇"],
            "room_number": row["房间号"],
            "capacity": _number(row["容纳人数"]),
            "equipment_summary": row["设备配置"],
            "open_hours": row["开放时间"],
            "status": row["开放状态"],
            "source": "07_实训室资源与排课台账_2026年8月.xlsx",
            "data_label": SOURCE_LABEL,
        }
        for row in rows
    ]


def _normalize_training_schedules(rows: list[dict[str, str]]) -> list[dict[str, Any]]:
    return [
        {
            "booking_id": row["排课编号"],
            "date": row["日期"],
            "start_time": row["开始时间"],
            "end_time": row["结束时间"],
            "room_id": row["实训室编号"],
            "class_name": row["班级/使用单位"],
            "activity": row["活动内容"],
            "owner_role": row["申请角色"],
            "status": row["安排状态"],
            "approval": row["审批状态"],
            "note": row["备注"],
            "source": "07_实训室资源与排课台账_2026年8月.xlsx",
            "data_label": SOURCE_LABEL,
        }
        for row in rows
    ]


def _normalize_training_equipment(rows: list[dict[str, str]]) -> list[dict[str, Any]]:
    return [
        {
            "equipment_id": row["设备编号"],
            "name": row["设备名称"],
            "model": row["设备型号"],
            "room_id": row["实训室编号"],
            "asset_status": row["设备状态"],
            "last_maintenance": row["最近维保"],
            "next_maintenance": row["下次维保"],
            "ticket_id": row["报修工单"],
            "issue": row["事项说明"],
            "ticket_status": row["处理状态"],
            "expected_complete_at": row["预计完成"],
            "source": "07_实训室资源与排课台账_2026年8月.xlsx",
            "data_label": SOURCE_LABEL,
        }
        for row in rows
    ]


def _normalize_training_safety_and_loans(rows: list[dict[str, str]]) -> list[dict[str, Any]]:
    return [
        {
            "record_id": row["记录编号"],
            "record_type": row["记录类型"],
            "date": row["日期"],
            "room_id": row["实训室编号"],
            "subject": row["事项"],
            "detail": row["事项说明"],
            "responsible_role": row["责任角色"],
            "status": row["状态"],
            "due_date": row["截止日期"],
            "action": row["处理建议"],
            "source": "07_实训室资源与排课台账_2026年8月.xlsx",
            "data_label": SOURCE_LABEL,
        }
        for row in rows
    ]


def _build_alerts(profiles: list[dict[str, Any]]) -> list[dict[str, Any]]:
    alerts: list[dict[str, Any]] = []
    for profile in profiles:
        reasons: list[str] = []
        if profile["attendance_absences_30d"] >= 3:
            reasons.append(f"近30天缺勤{profile['attendance_absences_30d']}次")
        if profile["failed_courses"]:
            reasons.append(f"不及格课程{profile['failed_courses']}门")
        if profile["training_log_missing"] or profile["training_issue_count"]:
            reasons.append("实训日志或操作记录存在待跟进事项")
        if "筛查建议" in profile["dorm_feedback"]:
            reasons.append("心理筛查建议开展关怀跟进（非诊断）")
        if not reasons:
            continue
        severity = "高关注" if len(reasons) >= 3 else "中关注"
        alerts.append(
            {
                "alert_id": f"ALT-202608-{profile['student_id']}",
                "student_id": profile["student_id"],
                "student_name": profile["name"],
                "class_name": profile["class_name"],
                "severity": severity,
                "reasons": reasons,
                "target_roles": ["班主任", "辅导员"],
                "status": "待核实",
                "generated_at": "2026-08-15 09:00",
                "data_label": SOURCE_LABEL,
                "note": "预警仅用于提示核实与帮扶分流，不构成处分、诊断或最终认定。",
            }
        )
    return alerts


def import_source_documents(store: PersistentStore, source_dir: Path) -> dict[str, Any]:
    """Read office source ledgers and upsert their normalized records into storage."""
    roster_file = source_dir / "01_学生基础名册_比赛演示.xlsx"
    attendance_file = source_dir / "02_考勤与学业汇总_2026年8月.xlsx"
    training_file = source_dir / "03_实训日志与操作记录_2026年8月.xlsx"
    screening_file = source_dir / "04_心理筛查结果_2026年秋季.xlsx"
    talk_file = source_dir / "05_辅导员谈心谈话记录_2026年8月.docx"
    task_file = source_dir / "06_学生帮扶任务清单_2026年8月.pdf"
    room_ledger_file = source_dir / "07_实训室资源与排课台账_2026年8月.xlsx"
    required = (
        roster_file, attendance_file, training_file, screening_file, talk_file,
        task_file, room_ledger_file,
    )
    missing = [path.name for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError(f"缺少待导入的演示源文件：{'、'.join(missing)}")

    roster = _rows_from_workbook(roster_file, "学生名册")
    attendance = _rows_from_workbook(attendance_file, "考勤与学业")
    training = _rows_from_workbook(training_file, "实训记录")
    screening = _rows_from_workbook(screening_file, "心理筛查")
    talks = _rows_from_docx(talk_file)
    tasks = _rows_from_pdf(task_file)
    training_rooms = _normalize_training_rooms(
        _rows_from_workbook(room_ledger_file, "实训室资源", "实训室编号")
    )
    training_schedules = _normalize_training_schedules(
        _rows_from_workbook(room_ledger_file, "场地排课", "排课编号")
    )
    training_equipment = _normalize_training_equipment(
        _rows_from_workbook(room_ledger_file, "设备状态", "设备编号")
    )
    training_safety_and_loans = _normalize_training_safety_and_loans(
        _rows_from_workbook(room_ledger_file, "安全与借用", "记录编号")
    )
    profiles = _build_profiles(roster, attendance, training, screening, talks)
    normalized_roster = _normalize_roster(roster)
    normalized_followups = _normalize_followups(talks)
    normalized_tasks = _normalize_tasks(tasks, profiles)
    alerts = _build_alerts(profiles)

    source_sets = (
        ("roster_students", normalized_roster, "student_id"),
        ("attendance_records", attendance, "记录编号"),
        ("training_operation_records", training, "记录编号"),
        ("psychological_screenings", screening, "记录编号"),
        ("followups", normalized_followups, "record_id"),
        ("support_tasks", normalized_tasks, "task_id"),
        ("student_profiles", profiles, "student_id"),
        ("risk_alerts", alerts, "alert_id"),
        ("training_rooms", training_rooms, "room_id"),
        ("training_room_schedules", training_schedules, "booking_id"),
        ("training_room_equipment", training_equipment, "equipment_id"),
        ("training_room_safety_and_loans", training_safety_and_loans, "record_id"),
    )
    for bucket, rows, id_field in source_sets:
        for row in rows:
            store.upsert_record(bucket, _text(row[id_field]), row)

    manifests = (
        _document_manifest(roster_file, len(roster), "学生基础名册"),
        _document_manifest(attendance_file, len(attendance), "考勤与学业台账"),
        _document_manifest(training_file, len(training), "实训日志与操作台账"),
        _document_manifest(screening_file, len(screening), "心理筛查台账"),
        _document_manifest(talk_file, len(talks), "谈心谈话记录"),
        _document_manifest(task_file, len(tasks), "帮扶任务清单"),
        _document_manifest(room_ledger_file, len(training_rooms) + len(training_schedules) + len(training_equipment) + len(training_safety_and_loans), "实训室资源与排课台账"),
    )
    # A source version supersedes the prior file manifest. Runtime follow-ups
    # and support tasks are kept in their own buckets, so this only prevents
    # stale file checksums from being shown as duplicate imported documents.
    store.replace_records("source_documents", list(manifests), "document_id")
    store.set_metadata("source_document_seed_version", SOURCE_VERSION)
    return {
        "source_version": SOURCE_VERSION,
        "source_label": SOURCE_LABEL,
        "documents": len(manifests),
        "roster_students": len(roster),
        "attendance_records": len(attendance),
        "training_operation_records": len(training),
        "psychological_screenings": len(screening),
        "followup_records": len(talks),
        "support_tasks": len(tasks),
        "risk_alerts": len(alerts),
        "training_rooms": len(training_rooms),
        "training_room_schedules": len(training_schedules),
        "training_room_equipment": len(training_equipment),
        "training_room_safety_and_loans": len(training_safety_and_loans),
    }


def next_followup_reminders(store: PersistentStore, days: int = 14) -> list[dict[str, Any]]:
    if days < 1 or days > 120:
        raise ValueError("提醒查询天数必须在1到120之间。")
    today = datetime.now().date()
    deadline = today.fromordinal(today.toordinal() + days)
    students = {item["student_id"]: item for item in store.list_records("student_profiles")}
    reminders: list[dict[str, Any]] = []
    for record in store.list_records("followups"):
        next_date = record.get("下次跟进日期") or record.get("next_followup_date") or ""
        if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", next_date):
            continue
        due_date = datetime.strptime(next_date, "%Y-%m-%d").date()
        if due_date > deadline:
            continue
        student = students.get(record.get("学号") or record.get("student_id"), {})
        reminders.append(
            {
                "reminder_id": f"REM-{record.get('记录编号') or record.get('record_id')}",
                "student_id": student.get("student_id", record.get("学号") or record.get("student_id")),
                "student_name": student.get("name", record.get("姓名", "")),
                "class_name": student.get("class_name", record.get("班级", "")),
                "owner": record.get("谈话人") or record.get("owner", "辅导员"),
                "next_followup_date": next_date,
                "status": "已逾期" if due_date < today else "即将到期",
                "next_action": record.get("下次行动") or record.get("next_action", "完成复访并更新谈话记录"),
                "data_label": SOURCE_LABEL,
            }
        )
    return sorted(reminders, key=lambda item: item["next_followup_date"])
